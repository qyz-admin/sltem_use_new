
import openpyxl
import docx

# 打开 Excel 文件并读取图片
wb = openpyxl.load_workbook('F:\需要用到的文件\退款模板\订单检索.xlsx')
ws = wb.active
img = openpyxl.drawing.image.Image(ws['A1'].value)

# 打开 Word 文件并查找替换内容
doc = docx.Document('example.docx')
for p in doc.paragraphs:
    if 'replace me' in p.text:
        # 替换内容为图片
        run = p.add_run()
        run.add_picture(img.filename, width=run._element.xpath('.//wp:extent')[0].get('cx'))

# 保存 Word 文件
doc.save('example_output.docx')
